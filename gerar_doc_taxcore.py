from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Criação do documento
document = Document()

# Título Principal
title = document.add_heading('Especificação de Engenharia: Arquitetura TaxCore BTP', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('Versão: 1.0 (Draft) | Foco: SPED Big 5 (ICMS/IPI, Contribuições, REINF, ECD, ECF)')
document.add_paragraph('_' * 70)  # Linha separadora

# --- SEÇÃO 1 ---
document.add_heading('1. Visão Geral da Arquitetura', level=1)
p1 = document.add_paragraph(
    'O objetivo é desacoplar a inteligência fiscal do ERP de origem (Clean Core Strategy). '
    'A solução opera em modelo Side-by-Side no SAP BTP, consumindo dados via APIs OData (modelo Pull) '
    'e processando obrigações de alto volume utilizando a capacidade in-memory do SAP HANA Cloud.'
)

document.add_heading('1.1. Pilares da Engenharia', level=2)
items_pilares = [
    'Ingestão via API ("Pull"): Substituição do SLT por chamadas OData agendadas e controladas por Delta Tokens.',
    'Camada de Abstração: O modelo de dados interno do TaxCore é único, independente se a origem é um S/4HANA Public Cloud ou Private Cloud.',
    'Processamento Push-down: Toda a lógica pesada de geração de arquivos (Layout 020, ECD, etc.) roda no nível do banco de dados (HANA AMDP).'
]
for item in items_pilares:
    document.add_paragraph(item, style='List Bullet')

# --- SEÇÃO 2 ---
document.add_heading('2. Camada de Ingestão de Dados (Data Ingestion Layer)', level=1)
document.add_paragraph('Tecnologia: SAP Integration Suite ou SAP Data Intelligence.\nProtocolo: OData v2/v4 (HTTPS com mTLS/OAuth).')

document.add_heading('2.1. Estratégia de Carga', level=2)
items_carga = [
    'Carga Inicial (Full Load): Particionamento por período (ex: Mês a Mês).',
    'Carga Incremental (Delta Load): Utilização obrigatória de filtros de timestamp (LastChangeDateTime) e paginação.',
    'Monitoramento: Dashboard de Health Check no BTP.'
]
for item in items_carga:
    document.add_paragraph(item, style='List Bullet')

document.add_heading('2.2. Mapa de APIs (OData) por Obrigação', level=2)

# Tabela de APIs
table = document.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Domínio'
hdr_cells[1].text = 'APIs Críticas (Standard)'
hdr_cells[2].text = 'Cobertura SPED'
hdr_cells[3].text = 'Observação Técnica'

data_apis = [
    ('Fiscal (NFe)', 'API_NOTA_FISCAL_SRV', 'ICMS/IPI, PIS/COFINS, REINF', 'Core. Extrai dados da J_1BNFDOC/LIN.'),
    ('Contábil', 'API_JOURNALENTRY_RETRIEVE_DIRECT_SRV', 'ECD, ECF', 'Alta performance para ler ACDOCA/BKPF.'),
    ('Estoque', 'API_MATERIAL_DOCUMENT_SRV', 'Bloco K e H, Crédito Estoque', 'Volume massivo. Requer particionamento.'),
    ('Produção', 'API_PRODUCTION_ORDER_2_SRV', 'Bloco K (K230/K235)', 'Leitura de ordens e empenhos.'),
    ('Mestres', 'API_BUSINESS_PARTNER / API_PRODUCT_SRV', 'Todos (0150, 0200)', 'Dados cadastrais básicos.'),
    ('Ativos', 'API_FIXEDASSET_SRV', 'CIAP (Bloco G), ECF', 'Cruzamento para crédito de ativo.')
]

for dominio, api, cob, obs in data_apis:
    row_cells = table.add_row().cells
    row_cells[0].text = dominio
    row_cells[1].text = api
    row_cells[2].text = cob
    row_cells[3].text = obs

# --- SEÇÃO 3 ---
document.add_heading('3. Camada de Persistência e Transformação (HANA Cloud)', level=1)
document.add_paragraph('Tecnologia: SAP HANA Cloud (Database as a Service).')

document.add_heading('3.1. Estrutura de Dados', level=2)
items_dados = [
    'Staging Area (Raw): Espelho exato das APIs (ex: STG_API_NOTA_FISCAL).',
    'Canonical Layer (Abstração TaxCore): Tabelas modeladas para a lógica fiscal brasileira, normalizadas.',
    'Reporting Layer (Virtual): Calculation Views que preparam os dados para o formato final.'
]
for item in items_dados:
    document.add_paragraph(item, style='List Bullet')

document.add_heading('3.2. Lifecycle Management', level=2)
document.add_paragraph('Hot Store (In-Memory) para dados recentes; Cold Store (Data Lake) para dados históricos.')

# --- SEÇÃO 4 ---
document.add_heading('4. Camada de Aplicação (Motor Fiscal)', level=1)
document.add_paragraph('Tecnologia: SAP BTP ABAP Environment ("Steampunk").')

items_app = [
    'Motor de Geração (AMDP): Classes ABAP que invocam SQLScript no HANA.',
    'Validador de Regras (BRF+): Configuração flexível.',
    'Job Scheduler: Gerenciamento de filas.'
]
for item in items_app:
    document.add_paragraph(item, style='List Bullet')

# --- SEÇÃO 5 ---
document.add_heading('5. Segurança e Conectividade', level=1)
document.add_paragraph('Tecnologia: SAP Cloud Connector (SCC) e SAP BTP Connectivity Service.')

items_sec = [
    'Private Cloud/On-Premise: Cloud Connector (Reverse Proxy).',
    'Public Cloud: HTTPS com OAuth 2.0.',
    'Criptografia: Dados em repouso e trânsito sempre criptografados.'
]
for item in items_sec:
    document.add_paragraph(item, style='List Bullet')

# --- SEÇÃO 6 ---
document.add_heading('6. Roadmap de Engenharia', level=1)
items_road = [
    '1. Proof of Concept (PoC): Conectar BTP ao S/4HANA e extrair API_NOTA_FISCAL.',
    '2. Definição do Modelo Canônico: Desenho das tabelas de abstração.',
    '3. Gap Analysis das APIs: Identificar campos faltantes (ex: Monofásico).'
]
for item in items_road:
    document.add_paragraph(item, style='List Number')

# Salvar o documento
file_name = 'Especificacao_Tecnica_TaxCore_BTP.docx'
document.save(file_name)

print(f"Documento '{file_name}' gerado com sucesso!")